"""Generate and run a simple CLI example for the Document Assembly Tool."""

import subprocess
from pathlib import Path

from docx import Document


def create_example_files(example_dir: Path) -> None:
    """Create a data source and template for the CLI example."""
    example_dir.mkdir(parents=True, exist_ok=True)

    # Create data source (Word table)
    data_doc = Document()
    data_doc.add_heading('Example Data', level=1)
    table = data_doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Field'
    table.rows[0].cells[1].text = 'Value'
    rows = [
        ('client_name', 'Acme Corp'),
        ('project_title', 'Enterprise Platform'),
        ('budget', '$150,000'),
    ]
    for i, (key, value) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = value

    data_path = example_dir / 'data.docx'
    data_doc.save(data_path)

    # Create template (placeholders)
    template = Document()
    template.add_heading('Project Proposal', level=1)

    p = template.add_paragraph()
    p.add_run('Client: ')
    run = p.add_run('{{client_name}}')
    run.bold = True

    template.add_paragraph('Project: {{project_title}}')
    template.add_paragraph('Budget: {{budget}}')
    template.add_paragraph('Thank you for your consideration.')

    template_path = example_dir / 'template.docx'
    template.save(template_path)

    return data_path, template_path


def run_example(example_dir: Path) -> None:
    data_path, template_path = create_example_files(example_dir)
    output_path = example_dir / 'output.docx'

    subprocess.run(
        [
            'python',
            '-m',
            'src.document_assembler',
            '-d',
            str(data_path),
            '-t',
            str(template_path),
            '-o',
            str(output_path),
        ],
        check=True,
    )

    print(f"Example output created: {output_path}")


if __name__ == '__main__':
    run_example(Path(__file__).parent)
