"""
Generate sample data and template files for testing.
"""

import csv
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK

def create_sample_data_word():
    """Create a Word document with a table of key-value pairs."""
    doc = Document()
    doc.add_heading('Sample Data', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Value'
    # Data rows
    rows = [
        ('client_name', 'Acme Corp'),
        ('project_title', 'Enterprise Platform'),
        ('budget', '$150,000')
    ]
    for i, (key, val) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = key
        # For value, we can add formatting if desired
        p = cells[1].paragraphs[0]
        p.add_run(val)
    fixture_dir = Path(__file__).parent / "fixtures"
    fixture_dir.mkdir(exist_ok=True)
    doc.save(str(fixture_dir / "sample_data.docx"))
    print("Created sample_data.docx")

def create_sample_data_csv():
    """Create a CSV file with one data row."""
    fixture_dir = Path(__file__).parent / "fixtures"
    path = fixture_dir / "sample_data.csv"
    with open(path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(['client_name', 'project_title', 'budget'])
        writer.writerow(['Acme Corp', 'Enterprise Platform', '$150,000'])
    print("Created sample_data.csv")

def create_sample_data_json():
    """Create a JSON file with an object."""
    data = {
        "client_name": "Acme Corp",
        "project_title": "Enterprise Platform",
        "budget": "$150,000"
    }
    fixture_dir = Path(__file__).parent / "fixtures"
    path = fixture_dir / "sample_data.json"
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)
    print("Created sample_data.json")

def create_sample_template():
    """Create a template document with placeholders."""
    doc = Document()
    doc.add_heading('Project Proposal', level=1)

    # Use a bold placeholder to verify inline replacement keeps formatting.
    p = doc.add_paragraph()
    p.add_run('Client: ')
    run = p.add_run('{{client_name}}')
    run.bold = True

    doc.add_paragraph('Project: {{project_title}}')
    doc.add_paragraph('Budget: {{budget}}')
    doc.add_paragraph('Thank you for your consideration.')

    fixture_dir = Path(__file__).parent / "fixtures"
    fixture_dir.mkdir(exist_ok=True)
    doc.save(str(fixture_dir / "sample_template.docx"))
    print("Created sample_template.docx")

if __name__ == "__main__":
    create_sample_data_word()
    create_sample_data_csv()
    create_sample_data_json()
    create_sample_template()
    print("All fixtures generated.")
