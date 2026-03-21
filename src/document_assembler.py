"""
Document Assembly Tool – Core Module

This module provides the DocumentAssembler class which extracts data from a structured
source (Word table, CSV, JSON) and inserts it into a template document at placeholder
locations, preserving all formatting. It supports customizable placeholder syntax
and multiple data source types.
"""

import argparse
import csv
import json
import logging
import re
import sys
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from docx import Document
from docx.text.paragraph import Paragraph
from opentelemetry import trace

# Configure module logger
logger = logging.getLogger(__name__)
tracer = trace.get_tracer(__name__)

from .cache.interface import TemplateCache

class DocumentAssembler:
    """
    Assembles a document by merging data from a structured source into a template.
    """

    # Default placeholder pattern (regex group name must match field name)
    DEFAULT_PLACEHOLDER_PATTERN = r'\{\{\s*(?P<field_name>\w+)\s*\}\}'

    def __init__(
        self,
        data_source: str,
        template_path: str,
        output_path: str,
        data_type: str = "word",
        placeholder_pattern: str = None,
        log_level: str = "INFO",
        cache: TemplateCache = None
    ):
        """
        Initialize the DocumentAssembler.
        """
        self.data_source = data_source
        self.template_path = template_path
        self.output_path = output_path
        self.data_type = data_type
        self.placeholder_pattern = placeholder_pattern or self.DEFAULT_PLACEHOLDER_PATTERN
        self.cache = cache

        # Set up logging
        logging.basicConfig(
            level=getattr(logging, log_level.upper()) if isinstance(log_level, str) else log_level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )

    def _get_template_hash(self) -> str:
        sha256 = hashlib.sha256()
        with open(self.template_path, 'rb') as f:
            for block in iter(lambda: f.read(4096), b''):
                sha256.update(block)
        return sha256.hexdigest()

    def _load_template(self) -> Document:
        with tracer.start_as_current_span("load_template"):
            if self.cache:
                template_hash = self._get_template_hash()
                cached_doc = self.cache.get(f"template:{template_hash}")
                if cached_doc:
                    logger.info("Template loaded from cache.")
                    from io import BytesIO
                    return Document(BytesIO(cached_doc))

            doc = Document(self.template_path)

            if self.cache:
                template_hash = self._get_template_hash()
                from io import BytesIO
                stream = BytesIO()
                doc.save(stream)
                self.cache.set(f"template:{template_hash}", stream.getvalue())
                logger.info("Template saved to cache.")

            return doc

    def _load_data_from_word(self) -> Dict[str, Any]:
        """Load data from a Word document table."""
        with tracer.start_as_current_span("load_data_word"):
            logger.info(f"Loading data from Word: {self.data_source}")
            doc = Document(self.data_source)
            if not doc.tables:
                raise ValueError("Word data source must contain at least one table.")

            table = doc.tables[0]
            data = {}
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                key = row.cells[0].text.strip()
                value_cell = row.cells[1]
                # Capture all paragraphs in the value cell to preserve formatting
                value_paras = []
                for para in value_cell.paragraphs:
                    runs_data = []
                    for run in para.runs:
                        runs_data.append({
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline
                        })
                    value_paras.append({
                        'style': para.style.name if para.style else None,
                        'runs': runs_data
                    })
                data[key] = value_paras
            logger.info(f"Loaded {len(data)} fields from Word.")
            return data

    def _load_data_from_csv(self) -> Dict[str, Any]:
        """Load data from CSV file."""
        with tracer.start_as_current_span("load_data_csv"):
            logger.info(f"Loading data from CSV: {self.data_source}")
            with open(self.data_source, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                if not rows:
                    raise ValueError("CSV file has no data rows.")
                data_row = rows[0]
            data = {}
            for key, value in data_row.items():
                data[key] = [{
                    'style': None,
                    'runs': [{'text': value, 'bold': False, 'italic': False, 'underline': False}]
                }]
            return data

    def _load_data_from_json(self) -> Dict[str, Any]:
        """Load data from JSON file."""
        with tracer.start_as_current_span("load_data_json"):
            logger.info(f"Loading data from JSON: {self.data_source}")
            with open(self.data_source, 'r', encoding='utf-8') as f:
                parsed = json.load(f)
            if isinstance(parsed, dict):
                data_dict = parsed
            elif isinstance(parsed, list) and parsed:
                data_dict = parsed[0]
            else:
                raise ValueError("JSON must be an object or non-empty array.")
            data = {}
            for key, value in data_dict.items():
                text = str(value)
                data[key] = [{
                    'style': None,
                    'runs': [{'text': text, 'bold': False, 'italic': False, 'underline': False}]
                }]
            return data

    def load_data(self) -> Dict[str, Any]:
        """Load data from the specified source."""
        if self.data_type == "word":
            return self._load_data_from_word()
        elif self.data_type == "csv":
            return self._load_data_from_csv()
        elif self.data_type == "json":
            return self._load_data_from_json()
        else:
            raise ValueError(f"Unsupported data type: {self.data_type}")

    def find_placeholders(self, doc: Document) -> List[tuple]:
        """
        Find all placeholders in the document.

        Returns:
            List of tuples (paragraph, match) where match is a regex match object.
        """
        placeholders = []
        pattern = re.compile(self.placeholder_pattern)
        for para in doc.paragraphs:
            matches = list(pattern.finditer(para.text))
            for match in matches:
                placeholders.append((para, match))
        return placeholders

    def insert_data(self, data: Dict[str, Any]) -> None:
        """
        Replace placeholders in the template with corresponding data.
        """
        with tracer.start_as_current_span("insert_data"):
            doc = self._load_template()
            pattern = re.compile(self.placeholder_pattern)

            for para in list(doc.paragraphs):
                if not pattern.search(para.text):
                    continue

                def _replacer(match):
                    field_name = match.group('field_name')
                    if field_name not in data:
                        return match.group(0)
                    values = ["".join(r['text'] for r in p['runs']) for p in data[field_name]]
                    return values[0] if values else ""

                for run in list(para.runs):
                    if pattern.search(run.text):
                        run.text = pattern.sub(_replacer, run.text)

            doc.save(self.output_path)
            logger.info(f"Assembled document saved to: {self.output_path}")

    def run(self) -> None:
        """Execute the full assembly process."""
        with tracer.start_as_current_span("assemble_document"):
            try:
                self.data = self.load_data()
                self.insert_data(self.data)
            except Exception as e:
                logger.critical(f"An unexpected error occurred: {e}", exc_info=True)
                sys.exit(1)

if __name__ == "__main__":
    pass
